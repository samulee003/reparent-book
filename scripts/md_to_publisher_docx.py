from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


TITLE = "養孩子，也重新好好養自己"
SUBTITLE = "心理師的覺察式育兒練習"


def set_run_font(run, ascii_font: str, east_asia_font: str, size_pt: int, bold=False, italic=False, color=None):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def add_toc(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = r'TOC \o "1-3" \h \z \u'

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "目錄將於 Word 開啟後自動更新。"
    placeholder.append(t)

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(placeholder)
    run._r.append(fld_char_end)


def apply_base_layout(doc: Document):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.3)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)

    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "新細明體")
    normal.font.size = Pt(12)
    p = normal.paragraph_format
    p.first_line_indent = Cm(0.85)
    p.line_spacing = 1.8
    p.space_before = Pt(0)
    p.space_after = Pt(0)

    for style_name, size, color in [
        ("Title", 24, "1F1F1F"),
        ("Subtitle", 14, "4A4A4A"),
        ("Heading 1", 18, "1F1F1F"),
        ("Heading 2", 14, "2F2F2F"),
        ("Heading 3", 12, "3C3C3C"),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微軟正黑體")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)

    styles["Title"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styles["Title"].paragraph_format.space_before = Pt(0)
    styles["Title"].paragraph_format.space_after = Pt(18)

    styles["Subtitle"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styles["Subtitle"].paragraph_format.space_after = Pt(10)

    styles["Heading 1"].paragraph_format.page_break_before = True
    styles["Heading 1"].paragraph_format.space_before = Pt(18)
    styles["Heading 1"].paragraph_format.space_after = Pt(14)

    styles["Heading 2"].paragraph_format.space_before = Pt(14)
    styles["Heading 2"].paragraph_format.space_after = Pt(8)

    styles["Heading 3"].paragraph_format.space_before = Pt(10)
    styles["Heading 3"].paragraph_format.space_after = Pt(4)

    if "Quote CN" not in [s.name for s in styles]:
        quote = styles.add_style("Quote CN", WD_STYLE_TYPE.PARAGRAPH)
        quote.base_style = styles["Normal"]
        quote.font.name = "Times New Roman"
        quote._element.rPr.rFonts.set(qn("w:eastAsia"), "新細明體")
        quote.font.size = Pt(11)
        quote.font.italic = True
        qpf = quote.paragraph_format
        qpf.left_indent = Cm(0.9)
        qpf.right_indent = Cm(0.9)
        qpf.first_line_indent = Cm(0)
        qpf.line_spacing = 1.6
        qpf.space_before = Pt(4)
        qpf.space_after = Pt(4)


def configure_header_footer(section):
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run(TITLE)
    set_run_font(hr, "Times New Roman", "微軟正黑體", 9, color="666666")

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(fp)


def parse_inline(paragraph, text: str, *, bold_default=False, italic_default=False):
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, "Times New Roman", "新細明體", 12, bold=True or bold_default, italic=italic_default)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, "Times New Roman", "新細明體", 12, bold=bold_default, italic=italic_default)


def collect_toc_entries(lines: list[str]) -> list[tuple[int, str]]:
    entries: list[tuple[int, str]] = []
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("# "):
            title = stripped[2:].strip()
            if title != "後記":
                entries.append((1, title))
            else:
                entries.append((1, "後記"))
        elif stripped.startswith("## "):
            title = stripped[3:].strip()
            if title == "本章小結":
                continue
            entries.append((2, title))
    return entries


def trim_duplicate_numbered_sections(lines: list[str]) -> list[str]:
    seen: set[tuple[int, int]] = set()
    trimmed: list[str] = []
    for line in lines:
        m = re.match(r"^##\s+(\d+)\.(\d+)", line.strip())
        if m:
            key = (int(m.group(1)), int(m.group(2)))
            if key in seen:
                break
            seen.add(key)
        trimmed.append(line)
    return trimmed


def load_book_lines(book_dir: Path) -> list[str]:
    ordered_files = [
        "foreword.md",
        "ch01.md",
        "ch02.md",
        "ch03.md",
        "ch04.md",
        "ch05.md",
        "ch06.md",
        "ch07.md",
        "ch08.md",
        "ch09.md",
        "ch10.md",
        "ch11.md",
        "ch12.md",
        "afterword.md",
    ]
    lines: list[str] = []
    for name in ordered_files:
        part = (book_dir / "02_contents" / name).read_text(encoding="utf-8").splitlines()
        part = trim_duplicate_numbered_sections(part)
        if lines and lines[-1].strip():
            lines.append("")
        lines.extend(part)
    return lines


def add_cover(doc: Document, toc_entries: list[tuple[int, str]]):
    p = doc.add_paragraph(style="Title")
    r = p.add_run(TITLE)
    set_run_font(r, "Times New Roman", "微軟正黑體", 24, bold=True)

    p2 = doc.add_paragraph(style="Subtitle")
    r2 = p2.add_run(SUBTITLE)
    set_run_font(r2, "Times New Roman", "新細明體", 14)

    for _ in range(6):
        doc.add_paragraph("")

    info = [
        "作者：Emily",
        "版本：完整書稿排版版",
        "日期：2026-04-09",
    ]
    for line in info:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        set_run_font(r, "Times New Roman", "新細明體", 11, color="555555")

    doc.add_page_break()

    toc_heading = doc.add_paragraph(style="Heading 1")
    r = toc_heading.add_run("目錄")
    set_run_font(r, "Times New Roman", "微軟正黑體", 18, bold=True)

    for level, text in toc_entries:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.first_line_indent = Cm(0)
        if level == 1:
            p.paragraph_format.left_indent = Cm(0)
            run = p.add_run(text)
            set_run_font(run, "Times New Roman", "微軟正黑體", 12, bold=True)
        else:
            p.paragraph_format.left_indent = Cm(0.8)
            run = p.add_run(text)
            set_run_font(run, "Times New Roman", "新細明體", 11, color="444444")
    doc.add_page_break()


def add_paragraph_with_style(doc: Document, text: str, style: str = "Normal"):
    p = doc.add_paragraph(style=style)
    if style == "Quote CN":
        parse_inline(p, text, italic_default=True)
    else:
        parse_inline(p, text)
    return p


def markdown_to_docx(md_path: Path, out_path: Path):
    raw_lines = load_book_lines(md_path.parent)

    toc_entries = collect_toc_entries(raw_lines)

    doc = Document()
    apply_base_layout(doc)
    configure_header_footer(doc.sections[0])
    add_cover(doc, toc_entries)

    paragraph_buffer: list[str] = []
    quote_buffer: list[str] = []

    def flush_paragraph():
        nonlocal paragraph_buffer
        if paragraph_buffer:
            text = " ".join(s.strip() for s in paragraph_buffer).strip()
            if text:
                add_paragraph_with_style(doc, text, "Normal")
            paragraph_buffer = []

    def flush_quote():
        nonlocal quote_buffer
        if quote_buffer:
            text = " ".join(s.strip() for s in quote_buffer).strip()
            if text:
                add_paragraph_with_style(doc, text, "Quote CN")
            quote_buffer = []

    for line in raw_lines:
        stripped = line.strip()

        if stripped == "---":
            flush_paragraph()
            flush_quote()
            doc.add_paragraph("")
            continue

        if not stripped:
            flush_paragraph()
            flush_quote()
            continue

        if stripped.startswith(">"):
            flush_paragraph()
            quote_buffer.append(stripped.lstrip(">").strip())
            continue

        if stripped.startswith("# "):
            flush_paragraph()
            flush_quote()
            p = doc.add_paragraph(style="Heading 1")
            r = p.add_run(stripped[2:].strip())
            set_run_font(r, "Times New Roman", "微軟正黑體", 18, bold=True)
            continue

        if stripped.startswith("## "):
            flush_paragraph()
            flush_quote()
            p = doc.add_paragraph(style="Heading 2")
            r = p.add_run(stripped[3:].strip())
            set_run_font(r, "Times New Roman", "微軟正黑體", 14, bold=True)
            continue

        if stripped.startswith("### "):
            flush_paragraph()
            flush_quote()
            p = doc.add_paragraph(style="Heading 3")
            r = p.add_run(stripped[4:].strip())
            set_run_font(r, "Times New Roman", "微軟正黑體", 12, bold=True)
            continue

        if stripped.startswith("- "):
            flush_paragraph()
            flush_quote()
            p = doc.add_paragraph(style="Normal")
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.hanging_indent = Cm(0.45)
            bullet = p.add_run("• ")
            set_run_font(bullet, "Times New Roman", "新細明體", 12)
            parse_inline(p, stripped[2:].strip())
            continue

        paragraph_buffer.append(stripped)

    flush_paragraph()
    flush_quote()

    # Ensure all later sections share header/footer
    for section in doc.sections[1:]:
        configure_header_footer(section)

    doc.save(out_path)


def main():
    if len(sys.argv) != 3:
        print("Usage: py -3 scripts/md_to_publisher_docx.py <input.md> <output.docx>")
        raise SystemExit(1)

    md_path = Path(sys.argv[1])
    out_path = Path(sys.argv[2])
    out_path.parent.mkdir(parents=True, exist_ok=True)
    markdown_to_docx(md_path, out_path)
    print(out_path.resolve())


if __name__ == "__main__":
    main()
