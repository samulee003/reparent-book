"""
generate_publisher_docx.py
出版社版 DOCX 生成腳本

《養孩子，也重新好好養自己：心理師的覺察式育兒練習》（全書以「安頓力」為核心框架）
讀取來源：book/plan-c/（方案 C 正式定稿版）
輸出格式：A4 / 25K 可調，微軟正黑體 + 新細明體雙字型，含封面、目錄、篇章頁、章節正文

用法:
    python scripts/generate_publisher_docx.py
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENTATION, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ─── 常數設定 ────────────────────────────────────────────────────────────────

REPO_ROOT = Path(__file__).resolve().parents[1]
BOOK_DIR = REPO_ROOT / "book" / "plan-c"

BOOK_TITLE = "養孩子，也重新好好養自己"
BOOK_SUBTITLE = "心理師的覺察式育兒練習"
AUTHOR = "鋅鋰師拔麻"
VERSION_LABEL = "出版社審閱版"
REVIEW_DATE = "2026-04-11"

OUTPUT = REPO_ROOT / "book" / f"養孩子，也重新好好養自己_出版社版_{REVIEW_DATE}.docx"

# 方案 C 四篇架構
PARTS = [
    {
        "number": "第一篇",
        "title": "先停下來，看見自己",
        "subtitle": "看見自己，先讓人回來",
        "files": ["ch01.md", "ch02.md", "ch03.md"],
    },
    {
        "number": "第二篇",
        "title": "舊傷從哪裡來",
        "subtitle": "理解與命名舊傷來源",
        "files": ["ch04.md", "ch05.md", "ch06.md", "ch07.md"],
    },
    {
        "number": "第三篇",
        "title": "重新養育自己",
        "subtitle": "安頓與修復自己",
        "files": ["ch08.md", "ch09.md"],
    },
    {
        "number": "第四篇",
        "title": "帶著覺察陪孩子長大",
        "subtitle": "把新的自己帶回教養現場，開始重寫",
        "files": ["ch10.md", "ch11.md", "ch12.md"],
    },
]

ORDERED_FILES_WITH_PART: list[tuple[str, int | None]] = [
    ("foreword.md", None),
    ("ch01.md", 0),
    ("ch02.md", 0),
    ("ch03.md", 0),
    ("ch04.md", 1),
    ("ch05.md", 1),
    ("ch06.md", 1),
    ("ch07.md", 1),
    ("ch08.md", 2),
    ("ch09.md", 2),
    ("ch10.md", 3),
    ("ch11.md", 3),
    ("ch12.md", 3),
    ("afterword.md", None),
]

# ─── 輔助函式 ────────────────────────────────────────────────────────────────


def set_run_font(
    run,
    ascii_font: str,
    east_asia_font: str,
    size_pt: int,
    bold: bool = False,
    italic: bool = False,
    color: str | None = None,
) -> None:
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph) -> None:
    """在段落中插入 PAGE 欄位（自動頁碼）。"""
    run = paragraph.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._r.append(instr)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


def configure_header_footer(section) -> None:
    section.different_first_page_header_footer = True
    # 首頁 header/footer 留空（封面不要頁碼和頁眉）
    first_hdr = section.first_page_header
    if not first_hdr.paragraphs:
        first_hdr.add_paragraph("")

    first_ftr = section.first_page_footer
    if not first_ftr.paragraphs:
        first_ftr.add_paragraph("")

    # 奇偶頁統一：書名頁眉 + 頁碼頁腳
    hdr = section.header
    hp = hdr.paragraphs[0]
    hp.clear()
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run(BOOK_TITLE)
    set_run_font(hr, "Times New Roman", "微軟正黑體", 9, color="888888")

    ftr = section.footer
    fp = ftr.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(fp)


def apply_base_layout(doc: Document) -> None:
    """設定基本頁面佈局與樣式。"""
    section = doc.sections[0]
    # A4 直向
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)

    styles = doc.styles

    # Normal 正文
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "新細明體")
    normal.font.size = Pt(12)
    nf = normal.paragraph_format
    nf.first_line_indent = Cm(0.84)  # 兩個全形字縮排
    nf.line_spacing = Pt(24)
    nf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    nf.space_before = Pt(0)
    nf.space_after = Pt(6)

    # 標題樣式
    for style_name, size_pt, color_hex, align in [
        ("Title", 28, "1A1A1A", WD_ALIGN_PARAGRAPH.CENTER),
        ("Subtitle", 16, "444444", WD_ALIGN_PARAGRAPH.CENTER),
        ("Heading 1", 20, "1A1A1A", WD_ALIGN_PARAGRAPH.LEFT),
        ("Heading 2", 15, "2C2C2C", WD_ALIGN_PARAGRAPH.LEFT),
        ("Heading 3", 12, "3C3C3C", WD_ALIGN_PARAGRAPH.LEFT),
    ]:
        sty = styles[style_name]
        sty.font.name = "Calibri"
        sty._element.rPr.rFonts.set(qn("w:eastAsia"), "微軟正黑體")
        sty.font.size = Pt(size_pt)
        sty.font.bold = True
        sty.font.color.rgb = RGBColor.from_string(color_hex)
        sty.paragraph_format.alignment = align

    styles["Title"].paragraph_format.space_before = Pt(0)
    styles["Title"].paragraph_format.space_after = Pt(12)

    styles["Subtitle"].paragraph_format.space_after = Pt(8)

    styles["Heading 1"].paragraph_format.page_break_before = True
    styles["Heading 1"].paragraph_format.space_before = Pt(24)
    styles["Heading 1"].paragraph_format.space_after = Pt(16)

    styles["Heading 2"].paragraph_format.page_break_before = False
    styles["Heading 2"].paragraph_format.space_before = Pt(16)
    styles["Heading 2"].paragraph_format.space_after = Pt(8)

    styles["Heading 3"].paragraph_format.space_before = Pt(12)
    styles["Heading 3"].paragraph_format.space_after = Pt(4)

    # 引用樣式
    if "Quote CN" not in [s.name for s in styles]:
        quote = styles.add_style("Quote CN", WD_STYLE_TYPE.PARAGRAPH)
        quote.base_style = styles["Normal"]
        quote.font.name = "Times New Roman"
        quote._element.rPr.rFonts.set(qn("w:eastAsia"), "新細明體")
        quote.font.size = Pt(11)
        quote.font.italic = True
        qf = quote.paragraph_format
        qf.left_indent = Cm(1.0)
        qf.right_indent = Cm(1.0)
        qf.first_line_indent = Cm(0)
        qf.line_spacing = Pt(22)
        qf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        qf.space_before = Pt(6)
        qf.space_after = Pt(6)

    # 列表樣式
    if "List CN" not in [s.name for s in styles]:
        lst = styles.add_style("List CN", WD_STYLE_TYPE.PARAGRAPH)
        lst.base_style = styles["Normal"]
        lst.font.name = "Times New Roman"
        lst._element.rPr.rFonts.set(qn("w:eastAsia"), "新細明體")
        lst.font.size = Pt(12)
        lf = lst.paragraph_format
        lf.first_line_indent = Cm(0)
        lf.left_indent = Cm(1.0)
        lf.line_spacing = Pt(24)
        lf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        lf.space_before = Pt(2)
        lf.space_after = Pt(2)


# ─── 行內格式 ────────────────────────────────────────────────────────────────


def parse_inline(paragraph, text: str, *, bold_default: bool = False, italic_default: bool = False, font_size: int = 12) -> None:
    """解析 **粗體**、*斜體*、`程式碼` 行內格式，並加入段落 runs。"""
    pattern = re.compile(r"(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)")
    pos = 0
    for m in pattern.finditer(text):
        # 前段一般文字
        if m.start() > pos:
            run = paragraph.add_run(text[pos : m.start()])
            set_run_font(run, "Times New Roman", "新細明體", font_size, bold=bold_default, italic=italic_default)
        if m.group(2):  # ***bold italic***
            run = paragraph.add_run(m.group(2))
            set_run_font(run, "Times New Roman", "新細明體", font_size, bold=True, italic=True)
        elif m.group(3):  # **bold**
            run = paragraph.add_run(m.group(3))
            set_run_font(run, "Times New Roman", "新細明體", font_size, bold=True)
        elif m.group(4):  # *italic*
            run = paragraph.add_run(m.group(4))
            set_run_font(run, "Times New Roman", "新細明體", font_size, italic=True)
        elif m.group(5):  # `code`
            run = paragraph.add_run(m.group(5))
            set_run_font(run, "Consolas", "新細明體", font_size - 1)
        pos = m.end()

    # 剩餘文字
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, "Times New Roman", "新細明體", font_size, bold=bold_default, italic=italic_default)


# ─── Markdown → DOCX 段落 ───────────────────────────────────────────────────


def process_lines(doc: Document, lines: list[str]) -> None:
    """逐行處理 Markdown 文本，將其轉換為 DOCX 段落。"""
    paragraph_buffer: list[str] = []
    quote_buffer: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if not paragraph_buffer:
            return
        text = " ".join(s.strip() for s in paragraph_buffer).strip()
        paragraph_buffer = []
        if not text:
            return
        p = doc.add_paragraph(style="Normal")
        parse_inline(p, text)

    def flush_quote() -> None:
        nonlocal quote_buffer
        if not quote_buffer:
            return
        text = " ".join(s.strip() for s in quote_buffer).strip()
        quote_buffer = []
        if not text:
            return
        p = doc.add_paragraph(style="Quote CN")
        parse_inline(p, text, italic_default=True, font_size=11)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 空行
        if not stripped:
            flush_paragraph()
            flush_quote()
            i += 1
            continue

        # 分隔線
        if re.match(r"^-{3,}$", stripped):
            flush_paragraph()
            flush_quote()
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.space_before = Pt(6)
            pf.space_after = Pt(6)
            # 用底線模擬分隔線
            from docx.oxml import OxmlElement as OXE
            pBdr = OXE("w:pBdr")
            btm = OXE("w:bottom")
            btm.set(qn("w:val"), "single")
            btm.set(qn("w:sz"), "4")
            btm.set(qn("w:space"), "1")
            btm.set(qn("w:color"), "CCCCCC")
            pBdr.append(btm)
            p._p.pPr.append(pBdr)
            i += 1
            continue

        # H1
        if stripped.startswith("# ") and not stripped.startswith("## "):
            flush_paragraph()
            flush_quote()
            title_text = stripped[2:].strip()
            p = doc.add_heading(title_text, level=1)
            p.style = doc.styles["Heading 1"]
            i += 1
            continue

        # H2
        if stripped.startswith("## ") and not stripped.startswith("### "):
            flush_paragraph()
            flush_quote()
            p = doc.add_heading(stripped[3:].strip(), level=2)
            i += 1
            continue

        # H3
        if stripped.startswith("### ") and not stripped.startswith("#### "):
            flush_paragraph()
            flush_quote()
            p = doc.add_heading(stripped[4:].strip(), level=3)
            i += 1
            continue

        # H4
        if stripped.startswith("#### "):
            flush_paragraph()
            flush_quote()
            p = doc.add_heading(stripped[5:].strip(), level=4)
            i += 1
            continue

        # 引用
        if stripped.startswith("> "):
            flush_paragraph()
            quote_buffer.append(stripped[2:])
            i += 1
            continue

        # 列表項
        list_m = re.match(r"^[-*•]\s+(.+)$", stripped)
        num_m = re.match(r"^\d+[.、)]\s+(.+)$", stripped)
        if list_m or num_m:
            flush_paragraph()
            flush_quote()
            item_text = (list_m or num_m).group(1)  # type: ignore[union-attr]
            bullet = "• " if list_m else ""
            p = doc.add_paragraph(style="List CN")
            p.paragraph_format.first_line_indent = Cm(0)
            if not num_m:
                r_bullet = p.add_run(bullet)
                set_run_font(r_bullet, "Times New Roman", "新細明體", 12)
            parse_inline(p, item_text)
            i += 1
            continue

        # 一般段落文字（累積）
        flush_quote()
        paragraph_buffer.append(stripped)
        i += 1

    flush_paragraph()
    flush_quote()


# ─── 封面 ──────────────────────────────────────────────────────────────────


def add_cover(doc: Document) -> None:
    """插入封面頁。"""
    # 大量空行推下封面內容
    for _ in range(8):
        doc.add_paragraph("")

    p_title = doc.add_paragraph(style="Title")
    r_title = p_title.add_run(BOOK_TITLE)
    set_run_font(r_title, "Calibri", "微軟正黑體", 28, bold=True, color="1A1A1A")
    p_title.paragraph_format.space_after = Pt(10)

    p_sub = doc.add_paragraph(style="Subtitle")
    r_sub = p_sub.add_run(BOOK_SUBTITLE)
    set_run_font(r_sub, "Calibri", "微軟正黑體", 16, color="444444")

    for _ in range(3):
        doc.add_paragraph("")

    p_author = doc.add_paragraph()
    p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_author = p_author.add_run(AUTHOR)
    set_run_font(r_author, "Calibri", "微軟正黑體", 14, color="333333")

    for _ in range(2):
        doc.add_paragraph("")

    p_ver = doc.add_paragraph()
    p_ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_ver = p_ver.add_run(f"{VERSION_LABEL}　{REVIEW_DATE}")
    set_run_font(r_ver, "Times New Roman", "新細明體", 10, color="888888")

    doc.add_page_break()


# ─── 目錄頁 ────────────────────────────────────────────────────────────────


def add_toc_page(doc: Document) -> None:
    """插入目錄標題與 Word 自動目錄欄位。"""
    p_head = doc.add_heading("目　錄", level=1)
    p_head.paragraph_format.page_break_before = False
    p_head.paragraph_format.space_before = Pt(8)
    p_head.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_toc = doc.add_paragraph()
    run = p_toc.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r' TOC \o "1-3" \h \z \u '

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "（在 Word 中按 Ctrl+A 再按 F9 可更新此目錄）"
    placeholder.append(t)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(placeholder)
    run._r.append(fld_end)

    doc.add_page_break()


# ─── 篇章頁 ──────────────────────────────────────────────────────────────────


def add_part_page(doc: Document, part: dict) -> None:
    """插入篇章頁（篇號 + 篇名 + 副標）。"""
    for _ in range(10):
        doc.add_paragraph("")

    p_num = doc.add_paragraph()
    p_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_num = p_num.add_run(part["number"])
    set_run_font(r_num, "Calibri", "微軟正黑體", 16, color="888888")

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_t = p_title.add_run(part["title"])
    set_run_font(r_t, "Calibri", "微軟正黑體", 24, bold=True, color="1A1A1A")

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_s = p_sub.add_run(part["subtitle"])
    set_run_font(r_s, "Times New Roman", "新細明體", 13, italic=True, color="555555")

    doc.add_page_break()


# ─── 主流程 ───────────────────────────────────────────────────────────────────


def build_document() -> None:
    doc = Document()
    apply_base_layout(doc)
    configure_header_footer(doc.sections[0])

    # 封面
    add_cover(doc)

    # 目錄
    add_toc_page(doc)

    # 前言
    foreword_lines = (BOOK_DIR / "foreword.md").read_text(encoding="utf-8").splitlines()
    process_lines(doc, foreword_lines)
    doc.add_page_break()

    # 各篇 & 各章
    current_part: int | None = None
    for file_name, part_idx in ORDERED_FILES_WITH_PART:
        if file_name in ("foreword.md", "afterword.md"):
            continue
        if part_idx != current_part:
            current_part = part_idx
            add_part_page(doc, PARTS[part_idx])  # type: ignore[index]

        chapter_lines = (BOOK_DIR / file_name).read_text(encoding="utf-8").splitlines()
        process_lines(doc, chapter_lines)
        doc.add_page_break()

    # 後記
    afterword_lines = (BOOK_DIR / "afterword.md").read_text(encoding="utf-8").splitlines()
    process_lines(doc, afterword_lines)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"✅ 已輸出：{OUTPUT}")


if __name__ == "__main__":
    build_document()
