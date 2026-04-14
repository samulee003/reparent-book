"""
generate_proposal_docx.py
將親子天下量身版出版企劃書（精簡版）輸出為 DOCX。
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

REPO_ROOT = Path(__file__).resolve().parents[1]
INPUT = REPO_ROOT / "book" / "親子天下量身版_出版企劃書_精簡版.md"
OUTPUT = REPO_ROOT / "book" / "親子天下量身版_出版企劃書_精簡版.docx"


def set_run_font(
    run,
    ascii_font: str = "Times New Roman",
    east_asia_font: str = "新細明體",
    size_pt: int = 12,
    bold: bool = False,
    italic: bool = False,
    color: str | None = None,
) -> None:
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def apply_base_layout(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "新細明體")
    normal.font.size = Pt(12)
    normal_pf = normal.paragraph_format
    normal_pf.line_spacing = Pt(22)
    normal_pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal_pf.space_before = Pt(0)
    normal_pf.space_after = Pt(6)

    for style_name, size_pt, color_hex, align in [
        ("Title", 22, "1A1A1A", WD_ALIGN_PARAGRAPH.CENTER),
        ("Heading 1", 18, "1A1A1A", WD_ALIGN_PARAGRAPH.LEFT),
        ("Heading 2", 14, "2C2C2C", WD_ALIGN_PARAGRAPH.LEFT),
        ("Heading 3", 12, "3C3C3C", WD_ALIGN_PARAGRAPH.LEFT),
    ]:
        sty = styles[style_name]
        sty.font.name = "Calibri"
        sty._element.rPr.rFonts.set(qn("w:eastAsia"), "微軟正黑體")
        sty.font.size = Pt(size_pt)
        sty.font.bold = True
        sty.font.color.rgb = RGBColor.from_string(color_hex)
        sty.paragraph_format.alignment = align

    styles["Title"].paragraph_format.space_after = Pt(14)
    styles["Heading 1"].paragraph_format.space_before = Pt(18)
    styles["Heading 1"].paragraph_format.space_after = Pt(10)
    styles["Heading 2"].paragraph_format.space_before = Pt(12)
    styles["Heading 2"].paragraph_format.space_after = Pt(6)
    styles["Heading 3"].paragraph_format.space_before = Pt(8)
    styles["Heading 3"].paragraph_format.space_after = Pt(4)

    if "List CN" not in [s.name for s in styles]:
        lst = styles.add_style("List CN", WD_STYLE_TYPE.PARAGRAPH)
        lst.base_style = styles["Normal"]
        lst.font.name = "Times New Roman"
        lst._element.rPr.rFonts.set(qn("w:eastAsia"), "新細明體")
        lst.font.size = Pt(12)
        lf = lst.paragraph_format
        lf.left_indent = Cm(0.8)
        lf.first_line_indent = Cm(0)
        lf.line_spacing = Pt(22)
        lf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        lf.space_before = Pt(2)
        lf.space_after = Pt(2)

    if "Quote CN" not in [s.name for s in styles]:
        quote = styles.add_style("Quote CN", WD_STYLE_TYPE.PARAGRAPH)
        quote.base_style = styles["Normal"]
        qf = quote.paragraph_format
        qf.left_indent = Cm(0.8)
        qf.right_indent = Cm(0.6)
        qf.first_line_indent = Cm(0)


def parse_inline(paragraph, text: str, font_size: int = 12) -> None:
    pattern = re.compile(r"(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos : m.start()])
            set_run_font(run, size_pt=font_size)
        if m.group(2):
            run = paragraph.add_run(m.group(2))
            set_run_font(run, size_pt=font_size, bold=True, italic=True)
        elif m.group(3):
            run = paragraph.add_run(m.group(3))
            set_run_font(run, size_pt=font_size, bold=True)
        elif m.group(4):
            run = paragraph.add_run(m.group(4))
            set_run_font(run, size_pt=font_size, italic=True)
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size_pt=font_size)


def parse_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def process_lines(doc: Document, lines: list[str]) -> None:
    paragraph_buffer: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if not paragraph_buffer:
            return
        text = " ".join(s.strip() for s in paragraph_buffer).strip()
        paragraph_buffer = []
        if not text:
            return
        p = doc.add_paragraph(style="Normal")
        parse_inline(p, text)

    i = 0
    while i < len(lines):
        stripped = lines[i].strip()

        if not stripped:
            flush_paragraph()
            i += 1
            continue

        if stripped.startswith("# "):
            flush_paragraph()
            p = doc.add_paragraph(style="Title")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped[2:].strip())
            set_run_font(run, "Calibri", "微軟正黑體", 22, bold=True)
            i += 1
            continue

        if stripped.startswith("## "):
            flush_paragraph()
            doc.add_heading(stripped[3:].strip(), level=1)
            i += 1
            continue

        if stripped.startswith("### "):
            flush_paragraph()
            doc.add_heading(stripped[4:].strip(), level=2)
            i += 1
            continue

        if stripped.startswith("> "):
            flush_paragraph()
            p = doc.add_paragraph(style="Quote CN")
            parse_inline(p, stripped[2:].strip(), font_size=11)
            i += 1
            continue

        list_m = re.match(r"^[-*]\s+(.+)$", stripped)
        num_m = re.match(r"^\d+[.)]\s+(.+)$", stripped)
        if list_m or num_m:
            flush_paragraph()
            p = doc.add_paragraph(style="List CN")
            bullet = "• " if list_m else ""
            if bullet:
                r = p.add_run(bullet)
                set_run_font(r, size_pt=12)
            parse_inline(p, (list_m or num_m).group(1))  # type: ignore[union-attr]
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and re.match(r"^\|(?:\s*:?-+:?\s*\|)+$", lines[i + 1].strip()):
            flush_paragraph()
            headers = parse_table_row(lines[i])
            rows: list[list[str]] = []
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(parse_table_row(lines[i]))
                i += 1

            table = doc.add_table(rows=1, cols=len(headers))
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            for idx, header in enumerate(headers):
                hdr_cells[idx].text = header
                for para in hdr_cells[idx].paragraphs:
                    for run in para.runs:
                        set_run_font(run, "Calibri", "微軟正黑體", 11, bold=True)

            for row in rows:
                cells = table.add_row().cells
                for idx, cell_text in enumerate(row):
                    cells[idx].text = cell_text
                    for para in cells[idx].paragraphs:
                        for run in para.runs:
                            set_run_font(run, size_pt=11)
            continue

        if re.match(r"^-{3,}$", stripped):
            flush_paragraph()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run("·  ·  ·")
            set_run_font(r, size_pt=12, color="AAAAAA")
            i += 1
            continue

        paragraph_buffer.append(stripped)
        i += 1

    flush_paragraph()


def build_docx() -> None:
    doc = Document()
    apply_base_layout(doc)
    lines = INPUT.read_text(encoding="utf-8").splitlines()
    process_lines(doc, lines)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"✅ 已輸出：{OUTPUT}")


if __name__ == "__main__":
    build_docx()
